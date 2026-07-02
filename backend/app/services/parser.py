import logging
from pathlib import Path
from docx import Document

# Настройка логирования (вывод в консоль)
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def extract_text_from_file(file_path: Path) -> list[str]:
    logger.info(f"extract_text_from_file called with: {file_path}")
    ext = file_path.suffix.lower()
    logger.info(f"File extension: {ext}")
    if ext == ".pdf":
        logger.info("Processing PDF file")
        return _extract_pdf(file_path)
    elif ext == ".docx":
        logger.info("Processing DOCX file")
        return [_extract_docx(file_path)]
    else:
        raise ValueError("Unsupported file format")

def _extract_pdf(file_path: Path) -> list[str]:
    logger.info("Entering _extract_pdf")
    pages = []
    try:
        from pypdf import PdfReader
        logger.info("pypdf imported successfully")
        with open(file_path, "rb") as f:
            logger.info("File opened for reading")
            reader = PdfReader(f)
            logger.info(f"PDF has {len(reader.pages)} pages")
            for page_num in range(len(reader.pages)):
                try:
                    text = reader.pages[page_num].extract_text() or ""
                    logger.info(f"Page {page_num+1} extracted, length: {len(text)}")
                except Exception as e:
                    logger.error(f"Error extracting page {page_num+1}: {e}")
                    text = ""
                pages.append(text)
        if not pages or all(p == "" for p in pages):
            logger.warning("No text extracted from PDF")
            raise ValueError("No text extracted from PDF")
    except Exception as e:
        logger.error(f"Failed to parse PDF: {e}")
        raise ValueError(f"Failed to parse PDF with pypdf: {e}")
    logger.info(f"PDF processing complete, {len(pages)} pages")
    return pages

def _extract_docx(file_path: Path) -> str:
    logger.info("Entering _extract_docx")
    try:
        logger.info("Attempting to open DOCX with python-docx")
        doc = Document(file_path)
        logger.info("DOCX opened successfully")
        full_text = "\n".join([para.text for para in doc.paragraphs])
        logger.info(f"Extracted paragraphs, total length: {len(full_text)}")
        if not full_text.strip():
            logger.info("No text in paragraphs, trying tables")
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text += cell.text + " "
            logger.info(f"Extracted tables, total length: {len(full_text)}")
        return full_text.strip()
    except Exception as e:
        logger.error(f"Failed to parse DOCX: {e}")
        raise ValueError(f"Failed to parse DOCX: {e}")